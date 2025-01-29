import streamlit as st
import openai
from openai import OpenAI
import pytesseract
import os
import tempfile
import cv2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Inches
from docx2pdf import convert
from io import BytesIO
from datetime import datetime
import numpy as np
import platform


# =======================
# 1) INITIALIZE SESSION STATE
# =======================
if "analysis_done" not in st.session_state:
    st.session_state["analysis_done"] = False
if "docx_stream" not in st.session_state:
    st.session_state["docx_stream"] = None
if "pdf_stream" not in st.session_state:
    st.session_state["pdf_stream"] = None
if "strengths" not in st.session_state:
    st.session_state["strengths"] = []
if "weaknesses" not in st.session_state:
    st.session_state["weaknesses"] = []
if "checklist" not in st.session_state:
    st.session_state["checklist"] = []

try:
    import streamlit as st
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
except (AttributeError, ModuleNotFoundError):
    import os
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY", ""))
    if not client.api_key:
        raise ValueError("OpenAI API key not found in environment variables or Streamlit secrets")

# Set Tesseract path based on OS
if platform.system() == 'Windows':
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
else:  # For Linux (Streamlit Cloud) and macOS
    pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

def extract_text_from_image(uploaded_file):
    """Extract text from image using Tesseract OCR with preprocessing."""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            # Write uploaded file content to temp file
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name
        
        # Load image using OpenCV
        img_cv = cv2.imread(tmp_path)
        
        # Convert to grayscale
        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
        
        # Apply thresholding to improve text clarity
        _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Apply morphological transformations to remove noise
        kernel = np.ones((1, 1), np.uint8)
        processed_img = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
        
        # Perform OCR
        text = pytesseract.image_to_string(processed_img).strip()
        
        # Clean up temporary file
        os.unlink(tmp_path)
        
        return text
    
    except Exception as e:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise RuntimeError(f"OCR failed: {str(e)}")

def build_prompt(post_text, rules):
    """
    Build the prompt for the GPT model based on the post text and rules.
    """
    rules_string = "\n".join([f"{rule['id']}: {rule['description']}" for rule in rules])
    prompt = (
        f"Analyze the following text for compliance with the specified rules:\n"
        f"Text: {post_text}\n\n"
        f"Rules:\n"
        f"{rules_string}\n\n"
        f"Identify if each rule is violated and provide a detailed comment for each."
        f"If a rule is violated, clearly specify why and flag it as 'Not Compliant.' "
        f"If no violation is found for a rule, mark it as 'Compliant.' "
        f"Be strict and flag exaggerations, misleading terms, or lack of risk warnings as violations."
    )
    return prompt

def create_compliance_table(analysis, rules):
    """
    Parse the analysis into strengths, weaknesses, and a compliance checklist.
    """
    strengths = []
    weaknesses = []
    checklist = []

    rule_lookup = {rule['id']: rule['description'] for rule in rules}

    for line in analysis.split("\n"):
        if ": Compliant" in line:
            rule_id, comment = line.split(": Compliant", 1)
            description = rule_lookup.get(rule_id.strip(), rule_id.strip())
            strengths.append({"Requirement": description, "Comments": comment.strip()})
            checklist.append({"Requirement": description, "Status": "✅ Compliant", "Comments": comment.strip()})
        elif ": Not Compliant" in line:
            rule_id, comment = line.split(": Not Compliant", 1)
            description = rule_lookup.get(rule_id.strip(), rule_id.strip())
            weaknesses.append({"Requirement": description, "Comments": comment.strip()})
            checklist.append({"Requirement": description, "Status": "❌ Not Compliant", "Comments": comment.strip()})

    # Ensure weaknesses match flagged content
    for weakness in weaknesses:
       ''' print(f"Flagged Non-Compliance: {weakness['Requirement']} - {weakness['Comments']}")'''

    return strengths, weaknesses, checklist


def generate_word_doc(
    post_text="",
    strengths=None,
    weaknesses=None,
    checklist=None,
    title_input="",
    references_input="",
    is_image=False,
    image_file=None,
    extracted_text=""
):
    if strengths is None: strengths = []
    if weaknesses is None: weaknesses = []
    if checklist is None: checklist = []

    doc = Document()

    # 1) Title
    if title_input.strip():
        p = doc.add_paragraph(title_input.strip())
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("")

    # 2) "Compliance Analysis Report"
    p2 = doc.add_paragraph("Compliance Analysis Report")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3) Date
       # Add current date in dd-mm-yy format
    current_date = datetime.now().strftime("%d-%m-%y")
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Date: {current_date}")
    date_run.bold = True
    date_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color

    # 4) Overview
    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "This document provides an analysis of the provided input "
        "for compliance with relevant rules and regulations."
    )

    # 5) Analyzed Input
    doc.add_heading("Analyzed Input", level=2)
    if is_image and image_file is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
            tmp_img.write(image_file.getvalue())
            tmp_path = tmp_img.name
        doc.add_picture(tmp_path, width=Inches(4))
        os.remove(tmp_path)

        if extracted_text.strip():
            doc.add_paragraph("Extracted Text:")
            doc.add_paragraph(extracted_text)
    else:
        doc.add_paragraph(post_text)

    # 6) Strengths
    doc.add_heading("Strengths", level=2)
    if strengths:
        for s in strengths:
            doc.add_paragraph(f"- {s['Requirement']}: {s['Comments']}")
    else:
        doc.add_paragraph("No strengths identified.")

    # 7) Weaknesses
    doc.add_heading("Weaknesses", level=2)
    if weaknesses:
        for w in weaknesses:
            doc.add_paragraph(f"- {w['Requirement']}: {w['Comments']}")
    else:
        doc.add_paragraph("No weaknesses identified.")

    # 8) Checklist
    doc.add_heading("Compliance Checklist", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Requirement"
    hdr_cells[1].text = "Status"
    hdr_cells[2].text = "Comments"

    for item in checklist:
        row_cells = table.add_row().cells
        row_cells[0].text = item["Requirement"]

        sp = row_cells[1].paragraphs[0]
        run = sp.add_run(item["Status"])
        if "✅" in item["Status"]:
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif "❌" in item["Status"]:
            run.font.color.rgb = RGBColor(255, 0, 0)

        row_cells[2].text = item["Comments"]

    # 9) References
    ref_list = [r.strip() for r in references_input.split("\n") if r.strip()]
    if ref_list:
        doc.add_heading("References", level=2)
        for i, ref in enumerate(ref_list, start=1):
            doc.add_paragraph(f"{i}. {ref}")

    docx_stream = BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    return docx_stream

def generate_pdf_from_docx(docx_stream):
    """
    Convert in-memory DOCX to PDF using docx2pdf.
    Requires Word on Win/Mac or LibreOffice on Linux.
    """
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
        tmp_docx.write(docx_stream.getvalue())
        tmp_docx_path = tmp_docx.name

    tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")
    convert(tmp_docx_path, tmp_pdf_path)

    pdf_stream = BytesIO()
    with open(tmp_pdf_path, "rb") as f:
        pdf_stream.write(f.read())
    pdf_stream.seek(0)

    os.remove(tmp_docx_path)
    os.remove(tmp_pdf_path)
    return pdf_stream


def analyze_compliance(post_text):
    """
    Analyze the given post text for compliance against predefined rules.
    :param post_text: The text of the blog or post to analyze.
    :return: A tuple of (strengths, weaknesses, checklist).
    """
    # Define rules and corresponding requirements
    rules = [
        {"id": "1.1", "rule": "Clarification of Expectations", "description": "Ensure expectations for financial promotions on social media are clarified."},
        {"id": "1.2", "rule": "Consumer Duty", "description": "Support retail customer understanding and decision-making."},
        {"id": "1.3", "rule": "Standalone Compliance", "description": "Each financial promotion must comply independently with regulations."},
        {"id": "1.4", "rule": "Balanced View", "description": "Present a balanced view of benefits and risks."},
        {"id": "1.5", "rule": "Specific Requirements", "description": "Include specific information prominently for high-risk investments."},
        {"id": "1.6", "rule": "Affiliate Marketer Responsibility", "description": "Ensure affiliate marketers comply with legal requirements."},
        {"id": "1.7", "rule": "Unauthorized Promotions", "description": "Unauthorized promotions without FCA approval are prohibited."},
        {"id": "1.9", "rule": "Influencers' Responsibility", "description": "Influencers must follow ASA rules and label paid promotions clearly."},
        {"id": "1.10", "rule": "Importance of Social Media", "description": "Social media can cause consumer harm due to complex financial products."},
        {"id": "1.12", "rule": "Guidance Context", "description": "Firms must comply with existing regulations and ensure clarity."},
        {"id": "2.1", "rule": "Definition of Financial Promotion", "description": "Communications inducing investment activity are financial promotions."},
        {"id": "2.1b", "rule": "Territorial Application", "description": "Financial promotion rules apply globally if they affect UK consumers."},
        {"id": "2.20", "rule": "Fair, Clear, Not Misleading", "description": "Promotions must avoid truncation or obscuring of key details."},
        {"id": "2.24", "rule": "Prominence of Information", "description": "Risk warnings must be clear and prominent."},
        {"id": "2.32", "rule": "Unsuitability of Certain Media", "description": "Complex products may not be suitable for social media promotions."},
        {"id": "2.38", "rule": "High-Risk Investments (HRIs) Restrictions", "description": "High-risk investments cannot be marketed to retail investors."},
        {"id": "2.41", "rule": "Prescribed Risk Warnings", "description": "Risk warnings must be clear, untruncated, and visible."},
        {"id": "3.1", "rule": "Consumer Duty in Strategies", "description": "Align marketing strategies with Consumer Duty obligations."},
        {"id": "3.3", "rule": "Target Market Tailoring", "description": "Tailor promotions to the characteristics of the target market."},
        {"id": "3.5", "rule": "Avoiding Behavioral Exploitation", "description": "Avoid exploiting behavioral biases, especially for vulnerable audiences."},
        {"id": "3.6", "rule": "Testing and Monitoring", "description": "Regularly monitor and test promotions for compliance."},
        {"id": "3.9", "rule": "Third-Party Sharing", "description": "Firms remain responsible for non-compliant third-party promotions."},
        {"id": "3.12", "rule": "Unsolicited Promotions", "description": "Ensure legal requirements for unsolicited promotions are met."},
        {"id": "4.2", "rule": "Influencer Business Models", "description": "Understand different influencer types and their risks."},
        {"id": "4.6", "rule": "Infographic Collaboration with ASA", "description": "Help influencers identify compliance risks."},
        {"id": "4.9", "rule": "Platform Responsibilities", "description": "Platforms must ensure hosted content complies with regulations."},
        {"id": "4.16", "rule": "In the Course of Business Criteria", "description": "Describe when influencers' actions constitute business activity."}
    ]

    # Build the prompt
    prompt = build_prompt(post_text, rules)


        # Query the GPT model
# Use the client to create a chat completion
        # Query the GPT model
        # Use the client to create a chat completion
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "user", "content": prompt}
            ],
        )
        # Parse the response
        analysis = response.choices[0].message.content
        # Create compliance table
        return create_compliance_table(analysis, rules)

    except openai.OpenAIError as e:

        print(f"API Error: {e}")
        return [], [], []
